"""
NCD Tracker - Term Sheet Generator
Renamed from term_sheet_generator.py.
Generates populated term sheets from template.
Works for both pipeline (draft) and closed deals.
"""

import re
from docx import Document
from docx.shared import RGBColor, Pt
from datetime import date, datetime
from pathlib import Path

import config
from utils import format_amount, format_percentage, format_date, format_tenor


class TermSheetGenerator:
    """Generates term sheets from a .docx template."""

    def __init__(self, template_path=None):
        self.template_path = Path(template_path) if template_path else config.TERM_SHEET_TEMPLATE

    # ══════════════════════════════════════════════
    #  PUBLIC API
    # ══════════════════════════════════════════════

    def generate_with_highlights(self, deal, output_path):
        """Generate term sheet (closed or draft) and highlight remaining placeholders."""
        doc_path = self.generate_term_sheet(deal, output_path)
        self.highlight_manual_fields(doc_path)
        return doc_path

    def generate_draft_for_pipeline(self, pipeline_deal, output_path):
        """
        Generate a DRAFT term sheet for a pipeline deal.
        Fields not yet known (ISIN, exact coupon) are left as [TBD] placeholders.
        """
        # Build a fake ClosedDeal-like object with available data
        from models import ClosedDeal
        draft = ClosedDeal(
            company_name    = pipeline_deal.company_name,
            instrument_type = pipeline_deal.instrument_type,
            issuer_type     = pipeline_deal.issuer_type,
            asset_class     = pipeline_deal.asset_class,
            issuance_size   = pipeline_deal.issuance_size,
            isin            = "[TBD – ISIN not yet assigned]",
            coupon          = 0.0,
            tenor           = 0,
            rating          = pipeline_deal.rating,
            security        = pipeline_deal.security,
            funding_date    = pipeline_deal.funding_date,
            maturity_date   = pipeline_deal.funding_date,  # TBD
        )
        return self.generate_with_highlights(draft, output_path)

    def generate_term_sheet(self, deal, output_path):
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        doc          = Document(self.template_path)
        replacements = self._build_replacements(deal)

        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        output_path = Path(output_path)
        doc.save(output_path)
        return output_path

    def highlight_manual_fields(self, doc_path):
        doc     = Document(doc_path)
        pattern = r'\[.*?\]'

        def _highlight(para):
            for run in para.runs:
                if re.search(pattern, run.text):
                    run.font.highlight_color = 7  # yellow

        for para in doc.paragraphs:
            _highlight(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _highlight(para)
        doc.save(doc_path)

    def batch_generate(self, closed_deals, output_folder):
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)
        paths = []
        for deal in closed_deals:
            safe = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_'
                           for c in deal.company_name)
            out  = output_folder / f"TermSheet_{safe}_{deal.isin}.docx"
            paths.append(self.generate_with_highlights(deal, out))
        return paths

    # ══════════════════════════════════════════════
    #  INTERNAL
    # ══════════════════════════════════════════════

    def _build_replacements(self, deal):
        coupon_str  = f"{deal.coupon:.2f}" if deal.coupon else "[TBD – Coupon Rate]"
        tenor_str   = str(deal.tenor) if deal.tenor else "[TBD – Tenor]"
        maturity_str = (format_date(deal.maturity_date)
                        if deal.maturity_date != deal.funding_date
                        else "[TBD – Maturity Date]")

        return {
            '[ISSUER_NAME]':      deal.company_name,
            '[ASSET_CLASS]':      deal.asset_class,
            '[ISSUER_TYPE]':      deal.issuer_type,
            '[ISIN_NUMBER]':      deal.isin,
            '[RATING]':           deal.rating,
            '[COUPON_RATE]':      coupon_str,
            '[AMOUNT]':           f"{deal.issuance_size:.2f}",
            '[AMOUNT_WORDS]':     self._number_to_words(deal.issuance_size),
            '[DATE]':             format_date(deal.funding_date),
            '[FUNDING_DATE]':     format_date(deal.funding_date),
            '[MATURITY_DATE]':    maturity_str,
            '[TENOR_MONTHS]':     tenor_str,
            '[TENOR_YEARS]':      f"{deal.tenor / 12:.1f}" if deal.tenor else "[TBD]",
            '[SECURITY_TYPE]':    deal.security,
            '[INSTRUMENT_TYPE]':  deal.instrument_type,
            '[TRUSTEE_NAME]':     'AUM Trustee Services Limited',
            '[LEGAL_COUNSEL]':    'Verdi Law',
            '[RATING_AGENCY]':    'IND Rating / ICRA / CARE',
            '[SPECIFY_USE_OF_FUNDS]': self._get_use_of_funds(deal.asset_class),
        }

    def _replace_in_paragraph(self, paragraph, replacements):
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

    def _number_to_words(self, number):
        ones  = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
        tens  = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
                 "Sixteen", "Seventeen", "Eighteen", "Nineteen"]

        def hundreds(n):
            if n == 0:   return ""
            if n < 10:   return ones[n]
            if n < 20:   return teens[n - 10]
            if n < 100:  return tens[n // 10] + (" " + ones[n % 10] if n % 10 else "")
            return ones[n // 100] + " Hundred" + (" " + hundreds(n % 100) if n % 100 else "")

        crores = int(number)
        if crores == 0:
            return "Zero Crores"
        result = hundreds(crores) + " Crore" + ("s" if crores != 1 else "")
        decimals = int((number - crores) * 100)
        if decimals:
            result += f" and {hundreds(decimals)} Paise"
        return result

    def _get_use_of_funds(self, asset_class):
        return {
            'NBFC':           'onward lending to customers as per its business model and general corporate purposes',
            'Housing Finance': 'providing housing loans and general corporate purposes',
            'MFI':            'providing microfinance loans and general corporate purposes',
            'Corporate':      'working capital, capex, and general corporate purposes',
            'EF':             'working capital, capex, and general corporate purposes',
            'FS':             'onward lending and general corporate purposes',
        }.get(asset_class, 'general corporate purposes')
