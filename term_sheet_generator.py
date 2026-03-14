"""
NCD Tracker - Term Sheet Generator
Generates populated term sheets from template for closed deals
"""

from docx import Document
from docx.shared import RGBColor, Pt
from datetime import date, datetime
from pathlib import Path
import re

import config
from data.data_models import ClosedDeal
from utils.formatters import format_amount, format_percentage, format_date, format_tenor

class TermSheetGenerator:
    """Generates term sheets from template"""
    
    def __init__(self, template_path=None):
        """Initialize generator with template"""
        self.template_path = Path(template_path) if template_path else config.TERM_SHEET_TEMPLATE
    
    def generate_term_sheet(self, closed_deal, output_path):
        """
        Generate term sheet for a closed deal
        
        Args:
            closed_deal (ClosedDeal): Closed deal data
            output_path (Path or str): Output file path
        
        Returns:
            Path: Path to generated document
        """
        # Load template
        doc = Document(self.template_path)
        
        # Build replacement dictionary
        replacements = self._build_replacements(closed_deal)
        
        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Save document
        output_path = Path(output_path)
        doc.save(output_path)
        
        return output_path
    
    def _build_replacements(self, deal):
        """
        Build dictionary of placeholder replacements
        
        Args:
            deal (ClosedDeal): Closed deal
        
        Returns:
            dict: Replacement mappings
        """
        replacements = {
            # Company info
            '[ISSUER_NAME]': deal.company_name,
            '[ASSET_CLASS]': deal.asset_class,
            
            # Instrument details
            '[ISIN_NUMBER]': deal.isin,
            '[RATING]': deal.rating,
            '[COUPON_RATE]': f"{deal.coupon:.2f}",
            
            # Amounts
            '[AMOUNT]': f"{deal.issuance_size:.2f}",
            '[AMOUNT_WORDS]': self._number_to_words(deal.issuance_size),
            
            # Dates
            '[DATE]': format_date(deal.funding_date),
            '[FUNDING_DATE]': format_date(deal.funding_date),
            '[MATURITY_DATE]': format_date(deal.maturity_date),
            
            # Tenor
            '[TENOR_MONTHS]': str(deal.tenor),
            '[TENOR_YEARS]': f"{deal.tenor / 12:.1f}",
            
            # Security type
            '[SECURITY_TYPE]': deal.security,
            
            # Instrument type flags
            '[INSTRUMENT_TYPE]': deal.instrument_type,
            
            # Default values for optional fields
            '[TRUSTEE_NAME]': 'AUM Trustee Services Limited',
            '[LEGAL_COUNSEL]': 'Verdi Law',
            '[RATING_AGENCY]': 'IND Rating / ICRA / CARE',
            
            # NBFC vs Corporate conditional text
            '[SPECIFY_USE_OF_FUNDS]': self._get_use_of_funds(deal.asset_class),
        }
        
        return replacements
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """
        Replace placeholders in a paragraph
        
        Args:
            paragraph: Document paragraph
            replacements (dict): Replacement mappings
        """
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    def _number_to_words(self, number):
        """
        Convert number to words (simplified, for Indian numbering)
        
        Args:
            number (float): Number in crores
        
        Returns:
            str: Number in words
        """
        # Simplified conversion - in production, use a library like num2words
        crores = int(number)
        decimals = int((number - crores) * 100)
        
        # Basic number to word conversion
        ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        
        def convert_hundreds(n):
            if n == 0:
                return ""
            elif n < 10:
                return ones[n]
            elif n < 20:
                return teens[n - 10]
            elif n < 100:
                return tens[n // 10] + (" " + ones[n % 10] if n % 10 > 0 else "")
            else:
                return ones[n // 100] + " Hundred" + (" " + convert_hundreds(n % 100) if n % 100 > 0 else "")
        
        if crores == 0:
            return "Zero Crores"
        
        result = convert_hundreds(crores) + " Crore"
        if crores != 1:
            result += "s"
        
        if decimals > 0:
            result += f" and {convert_hundreds(decimals)} Paise"
        
        return result
    
    def _get_use_of_funds(self, asset_class):
        """
        Get default use of funds text based on asset class
        
        Args:
            asset_class (str): Asset class
        
        Returns:
            str: Use of funds description
        """
        use_of_funds = {
            'NBFC': 'onward lending to its customers as per its business model and general corporate purposes',
            'Housing Finance': 'providing housing loans to customers and general corporate purposes',
            'MFI': 'providing microfinance loans to its customers and general corporate purposes',
            'Corporate': 'working capital requirements, capital expenditure, and general corporate purposes'
        }
        
        return use_of_funds.get(asset_class, 'general corporate purposes')
    
    def highlight_manual_fields(self, doc_path):
        """
        Highlight fields that require manual input in yellow
        
        Args:
            doc_path (Path): Document path
        """
        doc = Document(doc_path)
        
        # Fields that need manual input (still contain brackets)
        manual_fields_pattern = r'\[.*?\]'
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if re.search(manual_fields_pattern, run.text):
                    # Highlight in yellow
                    run.font.highlight_color = 7  # Yellow highlight
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if re.search(manual_fields_pattern, run.text):
                                run.font.highlight_color = 7
        
        doc.save(doc_path)
    
    def generate_with_highlights(self, closed_deal, output_path):
        """
        Generate term sheet and highlight manual fields
        
        Args:
            closed_deal (ClosedDeal): Closed deal data
            output_path (Path or str): Output file path
        
        Returns:
            Path: Path to generated document
        """
        # Generate term sheet
        doc_path = self.generate_term_sheet(closed_deal, output_path)
        
        # Highlight fields needing manual input
        self.highlight_manual_fields(doc_path)
        
        return doc_path
    
    def batch_generate(self, closed_deals, output_folder):
        """
        Generate term sheets for multiple deals
        
        Args:
            closed_deals (list): List of ClosedDeal objects
            output_folder (Path): Output folder path
        
        Returns:
            list: List of generated file paths
        """
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)
        
        generated_files = []
        
        for deal in closed_deals:
            # Create filename
            safe_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in deal.company_name)
            filename = f"TermSheet_{safe_name}_{deal.isin}.docx"
            output_path = output_folder / filename
            
            # Generate
            doc_path = self.generate_with_highlights(deal, output_path)
            generated_files.append(doc_path)
        
        return generated_files
